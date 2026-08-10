# -*- coding: utf-8 -*-
"""导出中心：Excel / CSV / Word / PDF / 编选报告。"""
from __future__ import annotations

import csv
import io
import os

STATUS_LABELS = {
    "pending": "待审核", "candidate": "候选", "shortlist": "初选",
    "final": "终选", "published": "最终出版", "rejected": "淘汰", "merged": "已合并",
}

FIELDS = [
    ("number", "编号"), ("title", "歌名"), ("uploader", "上传人"),
    ("category", "大类"), ("subcategory", "细类"),
    ("firstLine", "首句"), ("lyricist", "作者"),
    ("composer", "作曲"), ("translator", "译者"), ("tune", "曲调"), ("key", "调性"),
    ("meter", "格律"), ("themes", "主题"), ("scenarios", "场景"), ("musicTypes", "类型"),
    ("difficultyStars", "难度"), ("singabilityStars", "会众适唱"), ("rating", "评分"),
    ("status", "状态"), ("source", "来源"), ("comment", "备注"), ("updatedAt", "更新时间"),
]


def _val(song: dict, field: str):
    v = song.get(field, "")
    if field == "status":
        return STATUS_LABELS.get(v, v or "待审核")
    if isinstance(v, list):
        return " / ".join(str(x) for x in v)
    if v is None:
        return ""
    return str(v)


def rows_for(songs, fields=None):
    fields = fields or FIELDS
    return [[_val(s, f) for f, _ in fields] for s in songs]


def _style_header(ws, ncols, font_size=11):
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    thin = Side(style="thin", color="D9D9D9")
    for c in range(1, ncols + 1):
        cell = ws.cell(row=1, column=c)
        cell.font = Font(bold=True, color="FFFFFF", size=font_size, name="PingFang SC")
        cell.fill = PatternFill("solid", fgColor="4B6EAF")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = Border(bottom=thin)
    ws.freeze_panes = "A2"


def to_excel(songs, fields=None) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment
    fields = fields or FIELDS
    wb = Workbook()
    ws = wb.active
    ws.title = "曲目总表"
    ws.append([h for _, h in fields])
    _style_header(ws, len(fields))
    for s in songs:
        ws.append([_val(s, f) for f, _ in fields])
    widths = [10, 22, 24, 16, 16, 12, 16, 10, 10, 22, 20, 18, 10, 10, 8, 12, 16, 20, 20]
    for i, w in enumerate(widths[:len(fields)], start=1):
        ws.column_dimensions[chr(64 + i) if i <= 26 else "A"].width = w
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="center", wrap_text=True)
    ws.auto_filter.ref = ws.dimensions
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def to_csv(songs, fields=None) -> bytes:
    fields = fields or FIELDS
    out = io.StringIO()
    w = csv.writer(out)
    w.writerow([h for _, h in fields])
    for s in songs:
        w.writerow([_val(s, f) for f, _ in fields])
    return ("\ufeff" + out.getvalue()).encode("utf-8")


def to_docx(songs, title="赞美诗曲目表", subtitle="", fields=None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    fields = fields or FIELDS
    doc = Document()
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            r.font.size = Pt(10.5)
    table = doc.add_table(rows=1, cols=len(fields))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, (_, label) in enumerate(fields):
        hdr[i].text = label
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for s in songs:
        cells = table.add_row().cells
        for i, (f, _) in enumerate(fields):
            cells[i].text = _val(s, f)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _cjk_font():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    name = "HymnCJK"
    try:
        pdfmetrics.getFont(name)
    except Exception:
        candidates = [
            "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
            "/System/Library/Fonts/STHeiti Light.ttc",
        ]
        for path in candidates:
            if os.path.exists(path):
                try:
                    pdfmetrics.registerFont(TTFont(name, path))
                    break
                except Exception:
                    continue
    return name


def to_pdf(songs, title="赞美诗曲目表", subtitle="", fields=None) -> bytes:
    """每页单独生成再合并，保证中文与文本提取都正常。"""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from pypdf import PdfWriter, PdfReader
    fields = fields or FIELDS
    font = _cjk_font()
    pages = []
    page_w, page_h = A4
    margin = 40
    line_h = 16
    header_h = 46
    col_w = (page_w - 2 * margin) / len(fields)
    lines_per_page = int((page_h - 2 * margin - header_h) / line_h)

    def new_page():
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        c.setFont(font, 14)
        c.drawString(margin, page_h - margin - 14, title)
        c.setFont(font, 9)
        c.setFillColor(colors.HexColor("#666666"))
        c.drawString(margin, page_h - margin - 30, subtitle or "")
        c.setFillColor(colors.black)
        # header row
        y0 = page_h - margin - header_h
        c.setFillColor(colors.HexColor("#4B6EAF"))
        c.rect(margin, y0, page_w - 2 * margin, line_h, stroke=0, fill=1)
        c.setFillColor(colors.white)
        c.setFont(font, 8)
        for i, (_, label) in enumerate(fields):
            c.drawString(margin + i * col_w + 2, y0 + 5, label)
        c.setFillColor(colors.black)
        return c, buf

    c, buf = new_page()
    row_i = 0
    page_no = 1
    total_pages = max(1, (len(songs) + lines_per_page - 1) // lines_per_page) if songs else 1

    def finish_page():
        nonlocal page_no
        c.setFont(font, 8)
        c.setFillColor(colors.grey)
        c.drawString(page_w - margin - 60, margin - 20, f"第 {page_no}/{total_pages} 页")
        c.showPage()
        c.save()
        pages.append(buf.getvalue())
        page_no += 1

    for idx, s in enumerate(songs):
        if row_i >= lines_per_page:
            finish_page()
            c, buf = new_page()
            row_i = 0
        y = page_h - margin - header_h - line_h * (row_i + 1) + 4
        c.setFont(font, 7.5)
        for i, (f, _) in enumerate(fields):
            txt = _val(s, f)
            if len(txt) > 14:
                txt = txt[:14] + "…"
            c.drawString(margin + i * col_w + 2, y, txt)
        row_i += 1
    if row_i > 0 or not songs:
        finish_page()

    writer = PdfWriter()
    for p in pages:
        writer.append(PdfReader(io.BytesIO(p)))
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


# ---------------------------------------------------------------- 编选报告
def build_report(songs, title="赞美诗编选总表", subtitle="", extra_stats=None):
    """生成 (excel_bytes, word_bytes, pdf_bytes, csv_bytes)"""
    excel = to_excel(songs)
    csvb = to_csv(songs)
    word = to_docx(songs, title=title, subtitle=subtitle)
    pdf = to_pdf(songs, title=title, subtitle=subtitle)
    return excel, word, pdf, csvb
