# -*- coding: utf-8 -*-
"""导入解析：Excel / CSV / PDF / Word / 图片 / 音频 → 曲目字典列表。
识别策略：能提取什么就提取什么，重点保证「歌名 + 歌词」不丢。
PDF：文字层提取失败 → 逐页渲染图片 OCR（Vision → tesseract → AI）。
"""
from __future__ import annotations

import csv
import io
import json
import os
import pathlib
import re
import shutil
import subprocess
import tempfile
from difflib import SequenceMatcher

# ---------------------------------------------------------------- 列名自动匹配
COLUMN_ALIASES = {
    "title":      ["歌名", "歌曲名称", "歌曲名", "曲名", "名称", "标题", "赞美诗名称", "赞美诗", "诗名", "title", "歌名(中英)", "歌名中英"],
    "firstLine":  ["首句", "首句歌词", "第一句", "歌词首句", "首行", "first line", "firstline"],
    "lyricist":   ["作者", "作词", "词作者", "作词者", "作词人", "lyricist", "作词家", "词"],
    "composer":   ["作曲", "曲作者", "作曲者", "作曲人", "composer", "曲作者者"],
    "translator": ["译者", "翻译", "翻译者", "译词", "translator"],
    "tune":       ["曲调", "调名", "tune", "曲牌", "旋律"],
    "key":        ["调性", "调", "key", "原调", "定调"],
    "meter":      ["格律", "节拍", "拍号", "meter", "metre", "诗歌体"],
    "source":     ["来源", "出处", "歌本", "诗集", "source", "选本", "曲集"],
    "theme":      ["主题", "题材", "分类", "theme", "圣经主题", "类别"],
    "lyrics":     ["歌词", "歌词全文", "正文", "text", "lyrics", "full lyrics", "歌词内容"],
    "rating":     ["评分", "星级", "rating", "星"],
    "status":     ["状态", "status", "阶段"],
    "number":     ["编号", "序号", "number", "曲号", "诗编号"],
    "comment":    ["备注", "说明", "comment", "注释"],
    "difficulty": ["难度", "difficulty"],
    "uploader":   ["上传人", "上传人姓名", "收集人", "姓名", "uploader", "上报人"],
    "category":   ["大类", "category", "分类一", "圣诗分类"],
    "subcategory": ["细类", "subcategory", "分类二", "细目"],
}

_HEADER_CACHE = {}


def normalize_header(h: str) -> str:
    return re.sub(r"[\s\u3000·\.\-_]+", "", str(h or "")).lower()


def match_column(header: str):
    key = normalize_header(header)
    if key in _HEADER_CACHE:
        return _HEADER_CACHE[key]
    result = None
    for field, aliases in COLUMN_ALIASES.items():
        for a in aliases:
            if normalize_header(a) == key:
                result = field
                break
        if result:
            break
    if not result and key:
        best, best_r = None, 0.0
        for field, aliases in COLUMN_ALIASES.items():
            for a in aliases:
                r = SequenceMatcher(None, key, normalize_header(a)).ratio()
                if r > best_r:
                    best_r, best = r, field
        if best_r >= 0.62:
            result = best
    _HEADER_CACHE[key] = result
    return result


def _clean(value):
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def _map_rows(headers, rows):
    mapped = {}
    for i, h in enumerate(headers):
        f = match_column(h)
        if f:
            mapped[i] = f
    songs = []
    for row in rows:
        song = {}
        for i, f in mapped.items():
            v = _clean(row[i]) if i < len(row) else ""
            if v:
                song[f] = v
        if song:
            songs.append(song)
    return songs


# ---------------------------------------------------------------- Excel / CSV
def parse_excel_bytes(data: bytes, filename: str):
    name = filename.lower()
    if name.endswith(".csv"):
        return _parse_csv_bytes(data)
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return [], ["Excel 没有数据"]
    headers = [_clean(x) for x in rows[0]]
    songs = _map_rows(headers, rows[1:])
    return songs, []


def _try_decode_csv(data: bytes):
    for enc in ("utf-8-sig", "utf-8", "gb18030", "gbk", "big5"):
        try:
            return data.decode(enc), enc
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace"), "utf-8"


def _parse_csv_bytes(data: bytes):
    text, enc = _try_decode_csv(data)
    try:
        dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t")
    except Exception:
        dialect = csv.excel
    reader = list(csv.reader(io.StringIO(text), dialect))
    if not reader:
        return [], ["CSV 没有数据"]
    headers = [_clean(x) for x in reader[0]]
    songs = _map_rows(headers, reader[1:])
    return songs, []


# ---------------------------------------------------------------- 文本切分（宽松）
_SONG_NO_RE = re.compile(r"^\s*(?:第\s*)?(\d{1,4})\s*[、.．。)）\]\-\s:：]+\s*(\S.{0,40})$")
_HYMNBOOK_RE = re.compile(r"(?:赞美诗|诗歌|圣诗|颂赞)[^0-9]{0,12}?(?:第\s*)?(\d{1,4})\s*首?[：:\s]*(\S{2,40})")
_ANY_PUNCT = re.compile(r"[，。；：？！、…,.;:?!]")
_NOISE_RE = re.compile(r"^(第\s*\d{1,4}\s*(首|页)|page\s*\d+|\d{1,4}\s*页|目录|附录|序\s*$)", re.I)


def _is_noise_line(t: str) -> bool:
    t = t.strip()
    if not t:
        return True
    if re.fullmatch(r"[\d\s\-—.。·()（）]+", t):
        return True
    if _NOISE_RE.match(t):
        return True
    if re.search(r"^\d{1,3}\s*$", t):
        return True
    return False


def _is_title_candidate(t: str) -> bool:
    """歌名候选：短、无标点、含中英文文字。"""
    t = t.strip()
    if len(t) < 2 or len(t) > 20:
        return False
    if _ANY_PUNCT.search(t):
        return False
    return bool(re.search(r"[\u4e00-\u9fffA-Za-z]", t))


def _looks_like_title(t: str) -> bool:
    t = t.strip()
    if not t or len(t) > 42:
        return False
    if re.fullmatch(r"[\d\s\-—.。、()（）页pP]{1,10}", t):
        return False
    if not re.search(r"[\u4e00-\u9fffA-Za-z]", t):
        return False
    return True


def _entry_to_song(e: dict, source: str) -> dict:
    lyrics = e.get("lyrics", "").strip()
    title = e.get("title", "").strip()
    number = e.get("number", "")
    lines = [l for l in lyrics.splitlines() if l.strip()]
    if not title and lines and _is_title_candidate(lines[0]) and len(lines[0].strip()) <= 16:
        # 无编号时，把第一行短句当作歌名并从歌词中移除
        title = lines[0].strip()
        lines = lines[1:]
        lyrics = "\n".join(lines).strip()
    first = lines[0].strip() if lines else ""
    return {"title": title, "number": number, "lyrics": lyrics, "firstLine": first, "source": source}


def _segment_by_titles(lines, source):
    """无编号时的宽松切分：短行（无标点）视为歌名，其后为歌词。"""
    entries = []
    current = None
    for raw in lines:
        t = raw.strip()
        if not t or _is_noise_line(t):
            continue
        if _is_title_candidate(t):
            if current and current["lyrics"].strip():
                entries.append(current)
            current = {"number": "", "title": t, "lyrics": ""}
        else:
            if current is None:
                current = {"number": "", "title": "", "lyrics": ""}
            current["lyrics"] += (raw.strip() + "\n")
    if current and (current["lyrics"].strip() or current["title"]):
        entries.append(current)
    return [_entry_to_song(e, source) for e in entries]


def segment_hymn_text(text: str, source: str):
    """把整篇文本切分为 [{number,title,lyrics}]。优先编号模式，其次标题行宽松模式。"""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [l.rstrip() for l in text.split("\n")]
    entries = []
    current = None
    for raw in lines:
        t = raw.strip()
        if not t:
            continue
        m = _SONG_NO_RE.match(t) or _HYMNBOOK_RE.search(t)
        if m and _looks_like_title(m.group(2)):
            if current:
                entries.append(current)
            current = {"number": m.group(1),
                       "title": re.sub(r"\s+", " ", m.group(2)).strip(" \t-—。.，,"),
                       "lyrics": ""}
            continue
        if current is not None:
            current["lyrics"] += (raw + "\n")
    if current:
        entries.append(current)
    songs = [_entry_to_song(e, source) for e in entries if e["title"] or e["lyrics"].strip()]
    numbered = [s for s in songs if s.get("number")]
    if numbered:
        return numbered, []
    if songs:
        return songs, []
    return _segment_by_titles(lines, source), []


# ---------------------------------------------------------------- PDF（含扫描版 OCR 回退）
def render_pdf_pages(data: bytes, scale: float = 2.2):
    """把 PDF 每页渲染为 PIL 图像（供 OCR 识别扫描版）。"""
    import pypdfium2 as pdfium
    pdf = pdfium.PdfDocument(data)
    imgs = []
    for page in pdf:
        try:
            bitmap = page.render(scale=scale)
            imgs.append(bitmap.to_pil().convert("RGB"))
        except Exception:
            continue
    return imgs


def parse_pdf_bytes(data: bytes, filename: str, settings=None):
    warnings = []
    pages_text = []
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for pg in pdf.pages:
                try:
                    pages_text.append(pg.extract_text() or "")
                except Exception:
                    pages_text.append("")
    except Exception:
        pages_text = [""]
    full = "\n".join(pages_text)
    ocr_used = False
    if len(full.strip()) < 30:
        # 文字层不足 → 疑似扫描版：渲染成图片再 OCR
        try:
            imgs = render_pdf_pages(data)
        except Exception:
            imgs = []
        ocr_parts = []
        for im in imgs:
            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            tmp.close()
            try:
                im.save(tmp.name)
                text, engine, ai, lines = recognize_image_full2(tmp.name, settings)
                if text and text.strip():
                    ocr_parts.append((text.strip(), lines))
            finally:
                try:
                    os.unlink(tmp.name)
                except OSError:
                    pass
        if ocr_parts:
            full = "\n".join(ocr_parts)
            ocr_used = True
            warnings.append("扫描版 PDF：已通过 OCR 识别文字，请核对")
    if not full.strip():
        return [], ["PDF 无法识别文字（扫描版需 OCR：请安装 tesseract 或在设置中配置 AI 接口）"]
    if ocr_used:
        # 扫描版：每页支持一页多首切分
        songs = []
        for page_txt, page_lines in ocr_parts:
            if not page_txt.strip():
                continue
            multi = parse_ocr_lines_multi(page_lines, filename) if page_lines else []
            if multi:
                for m in multi:
                    songs.append({"title": m["title"], "firstLine": m["firstLine"],
                                  "lyrics": m["lyrics"], "source": filename, "flags": ["PDF-OCR"]})
            else:
                parsed = parse_ocr_plain_text(page_txt, filename)
                songs.append({"title": parsed["title"], "firstLine": parsed["firstLine"],
                              "lyrics": parsed["lyrics"], "source": filename,
                              "flags": ["PDF-OCR"]})
        return songs, warnings
    songs, seg_warn = segment_hymn_text(full, filename)
    warnings.extend(seg_warn)
    return songs, warnings


# ---------------------------------------------------------------- Word
def parse_docx_bytes(data: bytes, filename: str):
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = []
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    full = "\n".join(parts)
    if not full.strip():
        return [], ["Word 文档中没有可提取的文字"]
    songs, warnings = segment_hymn_text(full, filename)
    return songs, warnings


# ---------------------------------------------------------------- macOS Vision OCR
def ensure_ocr_tool():
    tools = pathlib.Path(__file__).resolve().parent / "tools"
    tool = tools / "ocr_tool"
    src = tools / "ocr_tool.swift"
    if tool.exists() and src.exists() and tool.stat().st_mtime >= src.stat().st_mtime:
        return str(tool)
    swiftc = shutil.which("swiftc")
    if not swiftc or not src.exists():
        return None
    cache = tools / ".cache"
    cache.mkdir(exist_ok=True)
    try:
        subprocess.run([swiftc, "-O", f"-Xcc", f"-fmodules-cache-path={cache}",
                        str(src), "-o", str(tool)], capture_output=True, timeout=240)
    except Exception:
        return None
    return str(tool) if tool.exists() else None


def ocr_engine_name():
    if ensure_ocr_tool():
        return "Vision"
    if _tessjs_ready():
        return "tesseract.js"
    if shutil.which("tesseract"):
        return "tesseract"
    return ""


def _tessjs_ready():
    node = _find_node()
    script = pathlib.Path(__file__).resolve().parent / "tools" / "ocr_tessera.js"
    td = pathlib.Path(__file__).resolve().parent / "data" / "tessdata" / "chi_sim.traineddata.gz"
    return bool(node and script.exists() and td.exists() and _find_node_modules())


def ocr_available():
    return bool(ocr_engine_name())


def ocr_image_vision(path):
    """macOS Vision（中文+英文）。返回 (text, lines) 或 None。"""
    tool = ensure_ocr_tool()
    if not tool:
        return None
    try:
        out = subprocess.run([tool, path], capture_output=True, timeout=150)
        if out.returncode != 0:
            return None
        lines = json.loads(out.stdout.decode("utf-8"))
        if not lines:
            return None
        text = "\n".join(l.get("text", "") for l in lines if l.get("text"))
        return (text, lines) if text.strip() else None
    except Exception:
        return None


def _find_node():
    import glob
    cands = [os.environ.get("NODE_BIN", ""),
             "/Users/macbook/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/bin/node"]
    cands += sorted(glob.glob(os.path.expanduser("~/.cache/codex-runtimes/*/dependencies/node/bin/node")), key=len)
    cands += [shutil.which("node") or ""]
    for c in cands:
        if c and os.path.exists(c):
            return c
    return None


def _find_node_modules():
    node = _find_node()
    if not node:
        return None
    cand = os.path.join(os.path.dirname(os.path.dirname(node)), "node_modules")
    return cand if os.path.isdir(cand) else None


def _preprocess_image_variant(path, scale_target=2200, contrast=None, binarize=False):
    """OCR 前预处理：EXIF 转正 + 缩放到指定大小 + 灰度 + 可选对比度/二值化。"""
    try:
        from PIL import Image, ImageOps, ImageEnhance
        im = Image.open(path)
        im = ImageOps.exif_transpose(im)  # 按 EXIF 转正
        im = im.convert("L")
        w, h = im.size
        m = max(w, h)
        if m > scale_target:
            sc = scale_target / m
            im = im.resize((int(w * sc), int(h * sc)), Image.LANCZOS)
        elif m < 900:
            sc = max(1.4, 900 / m)
            im = im.resize((int(w * sc), int(h * sc)), Image.LANCZOS)
        im = ImageOps.autocontrast(im)
        if contrast:
            im = ImageEnhance.Contrast(im).enhance(contrast)
        if binarize:
            im = im.point(lambda x: 0 if x < 175 else 255)
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tmp.close()
        im.save(tmp.name)
        return tmp.name
    except Exception:
        return path


def _preprocess_image(path):
    return _preprocess_image_variant(path, 2200, 1.5)


_OCR_PORT = int(os.environ.get("OCR_PORT", "8799"))
_ocr_server_proc = None
_ocr_server_checked = False


def _ocr_server_running():
    try:
        import urllib.request
        urllib.request.urlopen(f"http://127.0.0.1:{_OCR_PORT}/health", timeout=1)
        return True
    except Exception:
        return False


def _ensure_ocr_server():
    """确保持久 OCR 服务在跑（worker 常驻，批量识别快）。"""
    global _ocr_server_proc, _ocr_server_checked
    if _ocr_server_running():
        return _OCR_PORT
    node = _find_node()
    script = pathlib.Path(__file__).resolve().parent / "tools" / "ocr_server.js"
    mods = _find_node_modules()
    if not node or not script.exists() or not mods:
        return None
    env = dict(os.environ)
    env["NODE_PATH"] = mods
    try:
        _ocr_server_proc = subprocess.Popen([node, str(script)], env=env,
                                            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except Exception:
        return None
    for _ in range(40):
        if _ocr_server_running():
            return _OCR_PORT
        import time
        time.sleep(0.25)
    return None


def ocr_image_tesseractjs_batch(paths):
    """批量识别多张预处理后的图片（持久服务）。返回 [text,...] 或 None。"""
    port = _ensure_ocr_server()
    if not port:
        return None
    import urllib.request
    try:
        req = urllib.request.Request(f"http://127.0.0.1:{port}/ocr",
                                     data=json.dumps({"paths": paths}).encode("utf-8"),
                                     headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=900) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        texts = data.get("texts")
        return texts if isinstance(texts, list) else None
    except Exception:
        return None


def _run_tessjs_direct(pre, psm):
    node = _find_node()
    script = pathlib.Path(__file__).resolve().parent / "tools" / "ocr_tessera.js"
    mods = _find_node_modules()
    if not node or not script.exists() or not mods:
        return None
    env = dict(os.environ)
    env["NODE_PATH"] = mods
    try:
        out = subprocess.run([node, str(script), pre, psm], capture_output=True, timeout=300, env=env)
        return out.stdout.decode("utf-8", errors="replace").strip() or None
    except Exception:
        return None


def _run_tessjs_direct_lines(pre, psm):
    """运行 tesseract.js 并返回带坐标的行列表 [{text,x,y,w,h,conf}] 或 None。"""
    node = _find_node()
    script = pathlib.Path(__file__).resolve().parent / "tools" / "ocr_tessera.js"
    mods = _find_node_modules()
    if not node or not script.exists() or not mods:
        return None
    env = dict(os.environ)
    env["NODE_PATH"] = mods
    try:
        out = subprocess.run([node, str(script), pre, psm, "--json"], capture_output=True, timeout=300, env=env)
        if out.returncode != 0:
            return None
        lines = json.loads(out.stdout.decode("utf-8", errors="replace"))
        return lines if isinstance(lines, list) else None
    except Exception:
        return None


def parse_ocr_lines_multi(lines, source=""):
    """一页多首歌切分：大字标题 = 歌曲起点。
    规则：歌名 = 每首歌顶部大字（行高显著大于中位数的短行）；
    首行歌词 = 该歌曲谱（数字）行之后的第一行文字。
    返回 [{title, firstLine, lyrics}]；无法识别大字标题时返回 []。"""
    rows = [l for l in lines if l.get("text") and l["text"].strip()]
    if not rows:
        return []
    hs = sorted(l.get("h", 0) for l in rows)
    med = hs[len(hs) // 2] if hs else 0
    title_rows = []
    for l in rows:
        t = l["text"].strip()
        if (2 <= len(t) <= 14 and not re.search(r"[，。；：？！、…,.;:?!]", t)
                and l.get("h", 0) >= med * 1.15):
            title_rows.append(l)
    if not title_rows:
        return []
    # 按标题行切分
    songs = []
    cur = None
    for l in rows:
        t = l["text"].strip()
        if any(l is tr for tr in title_rows):
            if cur:
                songs.append(cur)
            cur = {"title": t, "seg": []}
        elif cur is not None:
            cur["seg"].append(t)
    if cur:
        songs.append(cur)

    result = []
    for song in songs:
        seg = song["seg"]
        lyric_lines = [l for l in seg if not _is_notation_line(l)]
        # 段内第一个曲谱行位置 → 之后的第一个文字行为首行
        first_nota = next((i for i, l in enumerate(seg) if _is_notation_line(l)), None)
        after = [l for i, l in enumerate(seg) if first_nota is not None and i > first_nota and not _is_notation_line(l)]
        first = after[0] if after else (lyric_lines[0] if lyric_lines else "")
        result.append({"title": song["title"], "firstLine": first, "lyrics": "\n".join(lyric_lines)})
    return result


def _pick_best_lines(lines_list):
    """多通道行结果中，取“歌曲数×歌名/首行完整度”最优。"""
    best, best_score = None, -1
    for lines in lines_list:
        if not lines:
            continue
        songs = parse_ocr_lines_multi(lines, "")
        if not songs:
            continue
        score = 0
        for sg in songs:
            score += (2 if sg["title"] else 0) + len(_CJK_RE.findall(sg["firstLine"] or ""))
        if score > best_score:
            best, best_score = lines, score
    return best


def ocr_image_tesseractjs_lines(path):
    """双通道识别，返回带坐标行（供一页多首切分）。"""
    if not (pathlib.Path(__file__).resolve().parent / "data" / "tessdata" / "chi_sim.traineddata.gz").exists():
        return None
    pre1 = _preprocess_image_variant(path, 2400, None)
    l1 = _run_tessjs_direct_lines(pre1, "4")
    try: os.unlink(pre1)
    except OSError: pass
    pre2 = _preprocess_image_variant(path, 2200, 1.5)
    l2 = _run_tessjs_direct_lines(pre2, "3")
    try: os.unlink(pre2)
    except OSError: pass
    return _pick_best_lines([l1, l2]) or (l1 or l2)


_TOC_NUM_RE = re.compile(r"^\s*(?:第\s*)?(\d{1,4})\s*[、.．)）\-:：\s]*\s*(.{1,30}?)(?:\s*[…·]{2,}\s*(\d{1,3}))?\s*$")
_TOC_SHOU_RE = re.compile(r"^\s*(?:第\s*)?(\d{1,3})\s*首[”\"']?\s*(.{1,30}?)\s*$")
_TOC_DOT_RE = re.compile(r"^\s*(.{2,26}?)[…·.。 ]{2,}(\d{1,3})\s*$")


def parse_toc_entries(text, source=""):
    """目录识别：从文本中提取 {number,title,page} 条目。至少 3 条才算目录。"""
    entries = []
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    for l in lines:
        if not l or re.fullmatch(r"[\d\s…·.。\-]+", l):
            continue
        m = _TOC_NUM_RE.match(l)
        if m:
            num = m.group(1)
            rest = m.group(2).strip()
            page = m.group(3) or ""
            if not page:
                mm = re.match(r"^(.*?)(\d{1,3})$", rest)
                if mm and len(mm.group(1)) >= 2:
                    rest, page = mm.group(1), mm.group(2)
            title = rest.strip(" […·.。、，,；;：:　]")
            if 2 <= len(title) <= 30:
                entries.append({"number": num, "title": title, "page": page})
            continue
        m2 = _TOC_DOT_RE.match(l)
        if m2:
            title = m2.group(1).strip(" […·.。、，,；;：:　]")
            if 2 <= len(title) <= 26:
                entries.append({"number": "", "title": title, "page": m2.group(2)})
            continue
        # “N首 + 歌名” 格式（赞美诗集目录常用：页码 + 71首 + 歌名）
        m3 = _TOC_SHOU_RE.match(l)
        if m3:
            num = m3.group(1)
            title = m3.group(2).strip(" ””\"'、，,；;:：…·.。 	")
            title = re.sub(r"\d{1,3}\s+\d{1,3}\s*首.*$", "", title)
            title = re.sub(r"\d{1,3}\s*首.*$", "", title)
            title = re.sub(r"\d{1,3}\s*$", "", title).strip(" ””\"'、，,；;:：…·.。 	")
            if 2 <= len(title) <= 30:
                entries.append({"number": num, "title": title, "page": ""})
    return entries if len(entries) >= 3 else []


def parse_toc_shou(text, source=""):
    """赞美诗集目录专用：行内含 “N首” → 取 N 为首数，N首之后文字为歌名。
    行尾混入的“页码 下一首号首”噪音一并清理。返回 [{number,title,page}]。"""
    entries = {}
    for l in text.splitlines():
        l = l.strip()
        m = re.search(r"(\d{1,3})\s*首[”\"']?", l)
        if not m:
            continue
        num = m.group(1)
        rest = l[m.end():].strip(" ””\"'、，,；;:：…·.。 	")
        rest = re.sub(r"\d{1,3}\s+\d{1,3}\s*首.*$", "", rest)
        rest = re.sub(r"\d{1,3}\s*首.*$", "", rest)
        rest = re.sub(r"\d{1,3}\s*$", "", rest)
        title = rest.strip(" ””\"'、，,；;:：…·.。 	")
        if 2 <= len(title) <= 30 and not re.fullmatch(r"[\d\s…·.。\-]+", title):
            if num not in entries or len(title) > len(entries[num]):
                entries[num] = title
    return [{"number": n, "title": t, "page": ""} for n, t in entries.items()]


def parse_ocr_text_multi(text, source=""):
    """一页多首切分（文字顺序版，不依赖坐标）：
    歌名 = 短行大字，且其下一行是数字曲谱行（如 51/3% 343…）→ 视为新歌起点。
    首行歌词 = 该首曲谱行之后的第一个文字行。返回 [{title, firstLine, lyrics}]。"""
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    titles = []
    for i, l in enumerate(lines):
        t = l
        if (2 <= len(t) <= 14 and not re.search(r"[，。；：？！、…,.;:?!]", t)
                and not _is_notation_line(t)):
            nxt = lines[i + 1] if i + 1 < len(lines) else ""
            if nxt and _is_notation_line(nxt) and re.search(r"\d", nxt):
                titles.append(i)
    if not titles:
        return []
    songs = []
    for k, idx in enumerate(titles):
        end = titles[k + 1] if k + 1 < len(titles) else len(lines)
        seg = lines[idx + 1:end]
        title = lines[idx]
        lyric_lines = [l for l in seg if not _is_notation_line(l)]
        first_nota = next((i for i, l in enumerate(seg) if _is_notation_line(l)), None)
        after = [l for i, l in enumerate(seg)
                 if first_nota is not None and i > first_nota and not _is_notation_line(l)]
        first = after[0] if after else (lyric_lines[0] if lyric_lines else "")
        songs.append({"title": title, "firstLine": first, "lyrics": "\n".join(lyric_lines)})
    return songs


def _run_tessjs_direct_words(pre, psm):
    """返回 {lines, words} 词级坐标数据。"""
    node = _find_node()
    script = pathlib.Path(__file__).resolve().parent / "tools" / "ocr_tessera.js"
    mods = _find_node_modules()
    if not node or not script.exists() or not mods:
        return None
    env = dict(os.environ)
    env["NODE_PATH"] = mods
    try:
        out = subprocess.run([node, str(script), pre, psm, "--json"], capture_output=True, timeout=300, env=env)
        if out.returncode != 0:
            return None
        data = json.loads(out.stdout.decode("utf-8", errors="replace"))
        return data if isinstance(data, dict) else None
    except Exception:
        return None


def reconstruct_rows(words):
    """按词级坐标把整页重组成行（不拆栏）：按 y 分组、行内按 x 排序。"""
    if not words:
        return []
    ws = sorted(words, key=lambda w: (w["y"], w["x"]))
    rows = []
    cur = []
    for w in ws:
        if cur and (w["y"] - cur[-1]["y"]) > max(6, cur[-1]["h"] * 0.7):
            rows.append(cur)
            cur = []
        cur.append(w)
    if cur:
        rows.append(cur)
    out = []
    for row in rows:
        row.sort(key=lambda w: w["x"])
        parts = []
        prev = None
        for w in row:
            t = w["text"]
            if prev and re.search(r"[A-Za-z0-9]$", prev) and re.match(r"^[A-Za-z0-9]", t):
                parts.append(" ")
            parts.append(t)
            prev = t
        out.append("".join(parts))
    return out


def reconstruct_columns(words):
    """按词级 x 坐标把页面分成左右栏，每栏重建为按上到下排列的行。"""
    if not words:
        return []
    centers = [w["x"] + w["w"] / 2 for w in words]
    gap, split = 0, None
    for a, b in zip(sorted(centers), sorted(centers)[1:]):
        if b - a > gap:
            gap, split = b - a, (a + b) / 2
    width = max(w["x"] + w["w"] for w in words) - min(w["x"] for w in words)
    if split is None or gap < width * 0.12:
        cols = [words]
    else:
        cols = [[w for w in words if w["x"] + w["w"] / 2 < split],
                [w for w in words if w["x"] + w["w"] / 2 >= split]]
    result = []
    for col in cols:
        col_sorted = sorted(col, key=lambda w: (w["y"], w["x"]))
        rows = []
        cur = []
        for w in col_sorted:
            if cur and (w["y"] - cur[-1]["y"]) > max(6, cur[-1]["h"] * 0.7):
                rows.append(cur)
                cur = []
            cur.append(w)
        if cur:
            rows.append(cur)
        col_lines = []
        for row in rows:
            row.sort(key=lambda w: w["x"])
            parts = []
            prev = None
            for w in row:
                t = w["text"]
                if prev and re.search(r"[A-Za-z0-9]$", prev) and re.match(r"^[A-Za-z0-9]", t):
                    parts.append(" ")
                parts.append(t)
                prev = t
            col_lines.append("".join(parts))
        result.append(col_lines)
    return result


def parse_column_lines(col_lines, source=""):
    """在“已按栏整理”的行列表里切多首：标题后跟数字曲谱行 = 新歌起点。"""
    lines = col_lines
    titles = []
    for i, l in enumerate(lines):
        if (2 <= len(l) <= 14 and not re.search(r"[，。；：？！、…,.;:?!]", l)
                and not _is_notation_line(l)):
            nxt = lines[i + 1] if i + 1 < len(lines) else ""
            if nxt and _is_notation_line(nxt) and re.search(r"\d", nxt):
                titles.append(i)
    if not titles:
        return []
    songs = []
    for k, idx in enumerate(titles):
        end = titles[k + 1] if k + 1 < len(titles) else len(lines)
        seg = lines[idx + 1:end]
        lyric_lines = [l for l in seg if not _is_notation_line(l)]
        first_nota = next((i for i, l in enumerate(seg) if _is_notation_line(l)), None)
        after = [l for i, l in enumerate(seg)
                 if first_nota is not None and i > first_nota and not _is_notation_line(l)]
        first = after[0] if after else (lyric_lines[0] if lyric_lines else "")
        songs.append({"title": lines[idx], "firstLine": first, "lyrics": "\n".join(lyric_lines)})
    return songs


def parse_ocr_columns_multi(words, source=""):
    """一页多栏多首：按栏重建 → 每栏按“歌名+数字曲谱行”切分。"""
    cols = reconstruct_columns(words)
    songs = []
    for col in cols:
        songs.extend(parse_column_lines(col, source))
    return songs


def ocr_image_tesseractjs_words(path):
    """三通道识别，返回最优的 {lines, words}。"""
    if not (pathlib.Path(__file__).resolve().parent / "data" / "tessdata" / "chi_sim.traineddata.gz").exists():
        return None
    cands = []
    for scale, contrast, binarize, psm in [(2400, None, False, "4"), (3000, None, True, "4"), (2200, 1.5, False, "3")]:
        pre = _preprocess_image_variant(path, scale, contrast, binarize=binarize)
        d = _run_tessjs_direct_words(pre, psm)
        try: os.unlink(pre)
        except OSError: pass
        if d:
            cands.append(d)
    best, best_score = None, -1
    for d in cands:
        songs = parse_ocr_columns_multi(d.get("words") or [], "")
        score = sum((2 if sg["title"] else 0) + len(_CJK_RE.findall(sg["firstLine"] or "")) for sg in songs)
        if score > best_score:
            best, best_score = d, score
    return best or (cands[0] if cands else None)


def _pick_best_ocr_text(texts):
    """多通道识别结果中，取“歌名+首行歌词+可切分首数”最完整的一个。"""
    best, best_score = None, -1
    for t in texts:
        if not t:
            continue
        p = parse_ocr_plain_text(t, "")
        score = (2 if p.get("title") else 0) + len(_CJK_RE.findall(p.get("firstLine") or ""))
        multi = parse_ocr_text_multi(t, "")
        score += len(multi) * 3  # 一页多首是重要信号
        if score > best_score:
            best, best_score = t, score
    return best


def ocr_image_tesseractjs_full(path):
    """三通道识别（一次返回 text + words 词级坐标），按“歌名/首行/多首切分”综合取最优。"""
    if not (pathlib.Path(__file__).resolve().parent / "data" / "tessdata" / "chi_sim.traineddata.gz").exists():
        return None, None
    cands = []
    for scale, contrast, binarize, psm in [(2400, None, False, "4"), (3000, None, True, "4"), (2200, 1.5, False, "3")]:
        pre = _preprocess_image_variant(path, scale, contrast, binarize=binarize)
        d = _run_tessjs_direct_words(pre, psm)
        try: os.unlink(pre)
        except OSError: pass
        if d:
            cands.append(d)
    best, best_score = None, -1
    for d in cands:
        t = (d.get("text") or "")
        score = 0
        if t:
            p = parse_ocr_plain_text(t, "")
            score += (2 if p["title"] else 0) + len(_CJK_RE.findall(p["firstLine"] or ""))
            score += len(parse_ocr_text_multi(t, "")) * 3
        words = d.get("words") or []
        if words:
            songs = parse_ocr_columns_multi(words, "")
            score += sum((2 if sg["title"] else 0) + len(_CJK_RE.findall(sg["firstLine"] or "")) for sg in songs) * 2
        if score > best_score:
            best, best_score = d, score
    if not best:
        return None, None
    return best.get("text") or None, best


def ocr_image_tesseractjs(path):
    text, _full = ocr_image_tesseractjs_full(path)
    return text


def ocr_image_tesseract(path):
    """tesseract（CPU，兼容无 GPU 环境；需安装 chi_sim）。"""
    if not shutil.which("tesseract"):
        return None
    try:
        out = subprocess.run(["tesseract", path, "stdout", "-l", "chi_sim+eng", "--psm", "6"],
                             capture_output=True, timeout=120)
        if out.returncode != 0:
            out = subprocess.run(["tesseract", path, "stdout", "-l", "eng", "--psm", "6"],
                                 capture_output=True, timeout=120)
        t = out.stdout.decode("utf-8", errors="replace").strip()
        return t or None
    except Exception:
        return None


# ---------------------------------------------------------------- AI 视觉识别（可选）
def ai_ocr_image(path, settings=None):
    """OpenAI 兼容视觉模型识别图片 → {title, firstLine, lyrics, number, note} 或 None。"""
    if not settings:
        return None
    key = (settings.get("openaiKey") or "").strip()
    base = (settings.get("openaiBase") or "https://api.openai.com/v1").rstrip("/")
    model = settings.get("openaiModel") or "gpt-4o-mini"
    if not key:
        return None
    import base64
    import urllib.request
    try:
        b64 = base64.b64encode(pathlib.Path(path).read_bytes()).decode("utf-8")
    except Exception:
        return None
    mime = "image/png"
    prompt = ("你是赞美诗资料整理助手。请识别这张赞美诗图片中的文字，只输出 JSON（不要输出其他内容）："
              '{"title":"歌名","firstLine":"歌词第一句","lyrics":"完整歌词，每句一行","number":"编号（没有则空字符串）"}')
    body = json.dumps({"model": model, "temperature": 0.1, "messages": [{"role": "user", "content": [
        {"type": "text", "text": prompt},
        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
    ]}]}).encode("utf-8")
    req = urllib.request.Request(base + "/chat/completions", data=body, headers={
        "Authorization": "Bearer " + key, "Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=90) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        obj = json.loads(data["choices"][0]["message"]["content"])
        return {"title": str(obj.get("title") or "").strip(),
                "firstLine": str(obj.get("firstLine") or "").strip(),
                "lyrics": str(obj.get("lyrics") or "").strip(),
                "number": str(obj.get("number") or "").strip(),
                "note": "由 AI 视觉识别，请核对后保存"}
    except Exception:
        return None


def recognize_image_full2(path, settings=None):
    """四级识别：Vision → tesseract.js → tesseract → AI。返回 (text, engine, ai_parsed, lines)。"""
    v = ocr_image_vision(path)
    if v:
        return v[0], "Vision", None, v[1]
    t, full = ocr_image_tesseractjs_full(path)
    if t:
        words = (full or {}).get("words") or []
        return t, "tesseract.js", None, {"words": words, "full": full}
    c = ocr_image_tesseract(path)
    if c:
        return c, "tesseract", None, None
    if settings and settings.get("openaiKey"):
        ai = ai_ocr_image(path, settings)
        if ai and (ai.get("lyrics") or ai.get("title")):
            txt = "\n".join(x for x in [ai.get("title"), ai.get("lyrics")] if x)
            return txt, "AI", ai, None
    return "", "", None, None


def recognize_image_full(path, settings=None):
    text, engine, ai, _lines = recognize_image_full2(path, settings)
    return text, engine, ai


def ocr_image(path):
    v = ocr_image_vision(path)
    if v:
        return v[0]
    return ocr_image_tesseract(path)


def parse_image_file(path, filename, settings=None):
    text, engine, ai, lines = recognize_image_full2(path, settings)
    if text and text.strip():
        if ai and ai.get("lyrics"):
            songs = [{"title": ai.get("title", ""), "firstLine": ai.get("firstLine", ""),
                      "lyrics": ai.get("lyrics", ""), "source": filename}]
        else:
            songs = []
            # ① 目录识别：优先“N首+歌名”格式，其次通用格式；≥3 条判为目录
            words = (lines or {}).get("words") if isinstance(lines, dict) else None
            rows = reconstruct_rows(words) if words else []
            rowtext = "\n".join(rows) if rows else text
            toc = parse_toc_shou(rowtext, filename)
            if len(toc) < 3:
                toc = parse_toc_entries(rowtext, filename)
            if len(toc) >= 3:
                songs = [{"title": e["title"], "number": e["number"],
                          "comment": ("页码 " + e["page"]) if e.get("page") else "来自目录",
                          "lyrics": "", "source": filename,
                          "ocrNote": "目录识别（无歌词，待补充）"} for e in toc]
            else:
                # ② 一页多首：词级分栏切分 → 文字顺序切分 → 单首
                words = (lines or {}).get("words") if isinstance(lines, dict) else None
                multi = parse_ocr_columns_multi(words, filename) if words else []
                if not multi:
                    multi = parse_ocr_text_multi(text, filename)
                if multi:
                    songs = [{"title": m["title"], "firstLine": m["firstLine"],
                              "lyrics": m["lyrics"], "source": filename} for m in multi]
                else:
                    parsed = parse_ocr_plain_text(text, filename)
                    songs = [{"title": parsed["title"], "firstLine": parsed["firstLine"],
                              "lyrics": parsed["lyrics"], "source": filename}]
        for sg in songs:
            sg.setdefault("flags", []).append("OCR识别")
        return songs, []
    return [], ["图片 OCR 未识别出文字（可安装 tesseract，或在设置中配置 AI 接口）"]


def parse_upload(filename: str, data: bytes, dest_dir: str, settings=None):
    name = os.path.basename(filename)
    safe = re.sub(r"[^\w.\-（）()\u4e00-\u9fff]+", "_", name)
    path = os.path.join(dest_dir, f"{os.urandom(4).hex()}_{safe}")
    with open(path, "wb") as f:
        f.write(data)
    ext = name.lower().rsplit(".", 1)[-1] if "." in name else ""
    attachment = {"name": name, "path": path, "ext": ext}

    if ext in ("xlsx", "xlsm", "xls"):
        songs, warns = parse_excel_bytes(data, name)
        kind = "excel"
    elif ext == "csv":
        songs, warns = _parse_csv_bytes(data)
        kind = "csv"
    elif ext == "pdf":
        songs, warns = parse_pdf_bytes(data, name, settings)
        kind = "pdf"
    elif ext in ("docx", "doc"):
        songs, warns = parse_docx_bytes(data, name)
        kind = "word"
    elif ext in ("jpg", "jpeg", "png", "bmp", "webp", "gif", "tif", "tiff"):
        songs, warns = parse_image_file(path, name, settings)
        kind = "image"
    elif ext in ("mp3", "wav", "m4a", "aac", "flac", "ogg", "wma"):
        songs, warns = parse_audio_file(path, name)
        kind = "audio"
    else:
        songs, warns = [], [f"暂不支持的文件类型 .{ext}"]
        kind = "other"

    return songs, warns, {"kind": kind, **attachment}


# ---------------------------------------------------------------- OCR 结果解析（拍照填表用）
def parse_song_from_ocr(lines, source=""):
    def is_noise(t):
        t = t.strip()
        if not t or len(t) < 2:
            return True
        if re.fullmatch(r"[\d\s\-.—·]{1,10}", t):
            return True
        if re.match(r"^(第\s*\d+\s*(首|页)|page\s*\d+|\d+\s*页)", t, re.I):
            return True
        return False

    rows = [l for l in lines if not is_noise(l.get("text", "")) and not _is_notation_line(l.get("text", ""))]
    if not rows:
        return {"title": "", "firstLine": "", "lyrics": "", "number": "", "note": "未识别出文字"}
    number = ""
    m = re.search(r"第\s*(\d{1,4})\s*首", " ".join(l.get("text", "") for l in rows))
    if m:
        number = m.group(1)
    hs = sorted(l.get("h", 0) for l in rows)
    med = hs[len(hs) // 2] if hs else 0
    title = ""
    for l in rows:
        t = l.get("text", "").strip()
        if 2 <= len(t) <= 14 and l.get("h", 0) >= med * 1.1 and not _ANY_PUNCT.search(t):
            title = t
            break
    if not title and rows:
        t0 = rows[0].get("text", "").strip()
        if 2 <= len(t0) <= 16:
            title = t0
    lyrics_lines = [l.get("text", "").strip() for l in rows
                    if l.get("text", "").strip() != title and not is_noise(l.get("text", ""))]
    lyrics = "\n".join(lyrics_lines)
    if not title:
        note = "未可靠识别歌名，请手动核对"
    elif not lyrics:
        note = "识别到歌名但未识别到歌词，请手动补充"
    else:
        note = "已自动提取歌名与歌词，请核对后保存"
    return {"title": title, "firstLine": (lyrics_lines[0] if lyrics_lines else ""),
            "lyrics": lyrics, "number": number, "note": note}


_CJK_RE = re.compile(r"[\u4e00-\u9fff]")


def _is_notation_line(line):
    """简谱/五线谱的曲谱行或乱码行：中文字符占比过低。"""
    t = (line or "").strip()
    if not t:
        return True
    total = len(re.sub(r"\s", "", t))
    if total == 0:
        return True
    cjk = len(_CJK_RE.findall(t))
    if cjk / total < 0.25:
        return True
    if re.match(r"^\s*\d+\s*[=＝]", t):  # 调号行：1=F
        return True
    return False


def parse_ocr_plain_text(text: str, source=""):
    """纯文本 OCR 结果 → {title, firstLine, lyrics, number, note}。
    歌谱规则：歌名 = 曲谱行之前的大字（短、无标点）；
    首行歌词 = 第一个曲谱（数字）行之后的第一个文字行。"""
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    # 找到第一个曲谱行（数字/符号行或调号行）
    first_nota = next((i for i, l in enumerate(lines) if _is_notation_line(l)), None)
    lyric_lines = [(i, l) for i, l in enumerate(lines) if not _is_notation_line(l)]

    # 首行歌词：第一个曲谱行之后的第一行文字
    after = [l for i, l in lyric_lines if first_nota is not None and i > first_nota]
    rest_all = [l for _, l in lyric_lines]
    first = after[0] if after else (rest_all[0] if rest_all else "")
    rest = after if after else rest_all

    # 歌名：曲谱行之前的第一行短文字（没有则取任意短行）
    before = [l for i, l in lyric_lines if first_nota is not None and i < first_nota]
    cands = before if before else rest_all
    title = ""
    for l in cands:
        if 2 <= len(l) <= 14 and not re.search(r"[，。；：？！、…,.;:?!]", l):
            title = l
            break
    if title and rest and rest[0] == title:
        rest = rest[1:]
    if first == title and rest:
        first = rest[0]
    lyrics = "\n".join(rest)
    if not title:
        note = "未识别到歌名，请手动填写（首行歌词已自动提取）" if first else "未识别出文字，请手动填写"
    elif not first:
        note = "识别到歌名但未识别到歌词，请手动补充"
    else:
        note = "已自动提取歌名与首行歌词，请核对后保存"
    return {"title": title, "firstLine": first, "lyrics": lyrics,
            "number": "", "note": note}


def parse_ocr_text(text: str, source=""):
    songs, _warns = segment_hymn_text(text, source)
    if songs and songs[0].get("number"):
        s = songs[0]
        return {"title": s["title"], "firstLine": s.get("firstLine", ""),
                "lyrics": s.get("lyrics", ""), "number": s.get("number", ""),
                "note": "已自动提取（识别到编号" + s.get("number", "") + "），请核对后保存"}
    return None
